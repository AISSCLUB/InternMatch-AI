"""
Unit & Integration Tests for End-to-End CV Pipeline Foundation.
Covers Document Parser, LLM Structured Extraction, Private Storage Download,
and Candidate Profile Write Repository Replacement semantics.
Uses TestingSessionLocal SQLite database fixture and zero real network calls.
"""

import io
from datetime import date
from unittest.mock import MagicMock
from uuid import uuid4

import pytest
from app.db.models import (
    Skill,
    StudentSkill,
)
from app.repositories.candidate_profile_write import (
    replace_candidate_profile_from_extraction,
)
from app.repositories.matching_data import MatchingDataRepository
from app.repositories.student_profile import StudentProfileRepository
from app.services.cv_parser import CVParsingError, extract_cv_text
from app.services.cv_profile_extraction import (
    ExtractedCandidateProfile,
    ExtractedEducation,
    ExtractedExperience,
    ExtractedPreferences,
    ExtractedProject,
    ExtractedSkill,
    extract_structured_candidate_profile,
    extract_structured_candidate_profile_multimodal,
)
from app.services.cv_storage import (
    CVStorageValidationError,
    download_candidate_cv,
)
from app.services.cv_validation import (
    CVValidationResult,
    CVValidationServiceError,
    InvalidCVDocumentError,
    validate_cv_document,
    validate_cv_document_multimodal,
)
from docx import Document

from tests.db import TestingSessionLocal


@pytest.fixture(autouse=True)
def _cleanup_cv_pipeline_created_skills():
    """
    Remove only Skill taxonomy rows created by the current CV pipeline test.

    The shared StaticPool database intentionally persists for the test process,
    so Gate 2.25 tests must not leak globally unique Skill names into later suites.
    """
    snapshot_db = TestingSessionLocal()
    try:
        initial_skill_ids = {skill_id for (skill_id,) in snapshot_db.query(Skill.id).all()}
    finally:
        snapshot_db.close()

    yield

    cleanup_db = TestingSessionLocal()
    try:
        current_skill_ids = {skill_id for (skill_id,) in cleanup_db.query(Skill.id).all()}
        created_skill_ids = current_skill_ids - initial_skill_ids

        if created_skill_ids:
            (
                cleanup_db.query(StudentSkill)
                .filter(StudentSkill.skill_id.in_(created_skill_ids))
                .delete(synchronize_session=False)
            )
            (
                cleanup_db.query(Skill)
                .filter(Skill.id.in_(created_skill_ids))
                .delete(synchronize_session=False)
            )

        cleanup_db.commit()
    except Exception:
        cleanup_db.rollback()
        raise
    finally:
        cleanup_db.close()


# ---------------------------------------------------------------------------
# 1. PARSER TESTS (1 - 6)
# ---------------------------------------------------------------------------


def test_pdf_parser_extracts_ordered_text(monkeypatch):
    """Test 1: PDF parser extracts ordered text across pages from in-memory content."""
    mock_page1 = MagicMock()
    mock_page1.extract_text.return_value = "Page 1: Jane Doe Software Engineer"
    mock_page2 = MagicMock()
    mock_page2.extract_text.return_value = "Page 2: Python, SQL, FastAPI"

    mock_reader = MagicMock()
    mock_reader.pages = [mock_page1, mock_page2]
    monkeypatch.setattr("app.services.cv_parser.PdfReader", lambda stream: mock_reader)

    pdf_bytes = b"%PDF-1.4 mock pdf content"
    result = extract_cv_text(storage_path="user_123/cv.pdf", content=pdf_bytes)

    assert "Page 1: Jane Doe Software Engineer" in result
    assert "Page 2: Python, SQL, FastAPI" in result
    assert result.index("Page 1") < result.index("Page 2")


def test_docx_parser_extracts_paragraph_order():
    """Test 2: DOCX parser extracts paragraph text in document order from real in-memory DOCX."""
    doc = Document()
    doc.add_paragraph("John Doe")
    doc.add_paragraph("Education: Computer Science")
    doc.add_paragraph("Experience: Backend Intern")
    buf = io.BytesIO()
    doc.save(buf)
    docx_bytes = buf.getvalue()

    result = extract_cv_text(storage_path="user_123/cv.docx", content=docx_bytes)

    assert "John Doe" in result
    assert "Education: Computer Science" in result
    assert "Experience: Backend Intern" in result
    assert result.index("John Doe") < result.index("Education") < result.index("Experience")


def test_unsupported_suffix_rejected():
    """Test 3: Non-PDF/DOCX file extension is rejected with CVParsingError."""
    with pytest.raises(CVParsingError, match="Unsupported file extension"):
        extract_cv_text(storage_path="user_123/cv.txt", content=b"some text content")


def test_empty_content_rejected():
    """Test 4: Empty byte payload is rejected with CVParsingError."""
    with pytest.raises(CVParsingError, match="cannot be empty"):
        extract_cv_text(storage_path="user_123/cv.pdf", content=b"")


def test_parser_failure_becomes_cv_parsing_error(monkeypatch):
    """Test 5: Low-level parser exceptions are converted to CVParsingError."""
    monkeypatch.setattr(
        "app.services.cv_parser.PdfReader",
        MagicMock(side_effect=Exception("Corrupt PDF byte stream")),
    )

    with pytest.raises(CVParsingError, match="Failed to parse PDF document"):
        extract_cv_text(storage_path="user_123/cv.pdf", content=b"corrupted bytes")


def test_no_useful_text_rejected(monkeypatch):
    """Test 6: Documents producing only empty/whitespace text are rejected with CVParsingError."""
    mock_page = MagicMock()
    mock_page.extract_text.return_value = "   \n\t  "
    mock_reader = MagicMock()
    mock_reader.pages = [mock_page]
    monkeypatch.setattr("app.services.cv_parser.PdfReader", lambda stream: mock_reader)

    with pytest.raises(CVParsingError, match="empty or whitespace-only"):
        extract_cv_text(storage_path="user_123/cv.pdf", content=b"%PDF content")


# ---------------------------------------------------------------------------
# 2. LLM EXTRACTION TESTS (7 - 13)
# ---------------------------------------------------------------------------


def _sample_extracted_profile():
    return ExtractedCandidateProfile(
        full_name="Samantha Ray",
        headline="AI Research Intern",
        skills=[
            ExtractedSkill(name="Python", proficiency_level="advanced"),
            ExtractedSkill(name="PyTorch", proficiency_level="intermediate"),
        ],
        education=[
            ExtractedEducation(
                institution="MIT",
                degree="B.S. Artificial Intelligence",
                start_year=2021,
                end_year=2025,
            )
        ],
        experience=[
            ExtractedExperience(
                company="OpenLab",
                role="Research Intern",
                description="Trained transformer models.",
                start_date=date(2023, 6, 1),
                end_date=date(2023, 8, 31),
            )
        ],
        projects=[
            ExtractedProject(
                title="Vision Transformer",
                tech_stack=["Python", "PyTorch"],
                description="Built vision model from scratch.",
            )
        ],
        preferences=ExtractedPreferences(
            work_types=["remote", "hybrid"],
            desired_locations=["Boston", "Remote"],
            target_roles=["AI Engineer", "ML Intern"],
        ),
    )


def test_gemini_client_constructed_at_call_time(monkeypatch):
    """Test 7: Gemini client is constructed dynamically during function call."""
    created_clients = []

    class MockGemini:
        def __init__(self, api_key):
            created_clients.append(api_key)
            self.models = MagicMock()
            mock_resp = MagicMock()
            mock_resp.text = _sample_extracted_profile().model_dump_json()
            self.models.generate_content.return_value = mock_resp

    monkeypatch.setattr("app.services.cv_profile_extraction.genai.Client", MockGemini)
    monkeypatch.setattr("app.core.config.settings.GEMINI_API_KEY", "gemini-test-key-12345")

    extract_structured_candidate_profile("Samantha Ray CV text")
    assert len(created_clients) == 1
    assert created_clients[0] == "gemini-test-key-12345"


def test_configured_llm_model_name_used(monkeypatch):
    """Test 8: Configured settings.LLM_MODEL_NAME is passed to Gemini generate_content call."""
    used_models = []

    class MockGemini:
        def __init__(self, api_key):
            self.models = MagicMock()

            def mock_generate(**kwargs):
                used_models.append(kwargs.get("model"))
                mock_resp = MagicMock()
                mock_resp.text = _sample_extracted_profile().model_dump_json()
                return mock_resp

            self.models.generate_content = mock_generate

    monkeypatch.setattr("app.services.cv_profile_extraction.genai.Client", MockGemini)
    monkeypatch.setattr("app.core.config.settings.GEMINI_API_KEY", "gemini-valid-key")
    monkeypatch.setattr("app.core.config.settings.LLM_MODEL_NAME", "gemini-3.5-flash-test")

    extract_structured_candidate_profile("Sample CV Text")
    assert used_models == ["gemini-3.5-flash-test"]


def test_blank_api_key_fails_before_provider_call(monkeypatch):
    """Test 9: Blank or placeholder GEMINI_API_KEY raises ValueError before client creation."""
    client_calls = []
    monkeypatch.setattr(
        "app.services.cv_profile_extraction.genai.Client",
        lambda api_key: client_calls.append(api_key),
    )

    monkeypatch.setattr("app.core.config.settings.GEMINI_API_KEY", "")
    with pytest.raises(ValueError, match="GEMINI_API_KEY"):
        extract_structured_candidate_profile("CV text")

    monkeypatch.setattr("app.core.config.settings.GEMINI_API_KEY", "gemini-placeholder-key-value")
    with pytest.raises(ValueError, match="GEMINI_API_KEY"):
        extract_structured_candidate_profile("CV text")

    assert client_calls == []


def test_structured_parsed_profile_returned(monkeypatch):
    """Test 10: Valid Gemini response returns ExtractedCandidateProfile instance."""
    expected_profile = _sample_extracted_profile()

    class MockGemini:
        def __init__(self, api_key):
            self.models = MagicMock()
            mock_resp = MagicMock()
            mock_resp.text = expected_profile.model_dump_json()
            self.models.generate_content.return_value = mock_resp

    monkeypatch.setattr("app.services.cv_profile_extraction.genai.Client", MockGemini)
    monkeypatch.setattr("app.core.config.settings.GEMINI_API_KEY", "gemini-valid-key")

    result = extract_structured_candidate_profile("Samantha Ray CV text")
    assert isinstance(result, ExtractedCandidateProfile)
    assert result.full_name == "Samantha Ray"
    assert len(result.skills) == 2
    assert result.skills[0].name == "Python"
    assert len(result.education) == 1
    assert result.education[0].institution == "MIT"


def test_empty_response_rejected(monkeypatch):
    """Test 11: Empty model response text raises ValueError."""

    class MockGemini:
        def __init__(self, api_key):
            self.models = MagicMock()
            mock_resp = MagicMock()
            mock_resp.text = ""
            self.models.generate_content.return_value = mock_resp

    monkeypatch.setattr("app.services.cv_profile_extraction.genai.Client", MockGemini)
    monkeypatch.setattr("app.core.config.settings.GEMINI_API_KEY", "gemini-valid-key")

    with pytest.raises(ValueError, match="Model returned unparseable or empty structured output"):
        extract_structured_candidate_profile("CV text")


def test_llm_provider_exception_propagates(monkeypatch):
    """Test 12: Genuine Gemini provider exception propagates unchanged."""

    class MockGemini:
        def __init__(self, api_key):
            self.models = MagicMock()
            self.models.generate_content.side_effect = RuntimeError(
                "Gemini 500 internal server error"
            )

    monkeypatch.setattr("app.services.cv_profile_extraction.genai.Client", MockGemini)
    monkeypatch.setattr("app.core.config.settings.GEMINI_API_KEY", "gemini-valid-key")

    with pytest.raises(RuntimeError, match="Gemini 500 internal server error"):
        extract_structured_candidate_profile("CV text")


def test_import_cv_profile_extraction_performs_no_work(monkeypatch):
    """Test 13: Importing module performs zero network or client creation."""
    calls = []
    monkeypatch.setattr(
        "app.services.cv_profile_extraction.genai.Client",
        lambda api_key: calls.append(api_key),
    )

    import app.services.cv_profile_extraction  # noqa: F401

    assert calls == []


# ---------------------------------------------------------------------------
# 3. STORAGE DOWNLOAD BOUNDARY TESTS (14 - 19)
# ---------------------------------------------------------------------------


def make_mock_supabase_client(download_bytes=b"%PDF content"):
    mock_bucket = MagicMock()
    mock_bucket.download.return_value = download_bytes
    mock_storage = MagicMock()
    mock_storage.from_.return_value = mock_bucket
    mock_client = MagicMock()
    mock_client.storage = mock_storage
    return mock_client, mock_bucket


def test_storage_download_uses_configured_bucket(monkeypatch):
    """Test 14: download_candidate_cv queries the configured settings.CV_STORAGE_BUCKET."""
    mock_client, mock_bucket = make_mock_supabase_client()
    monkeypatch.setattr("app.services.cv_storage.create_client", lambda url, key: mock_client)
    monkeypatch.setattr("app.core.config.settings.SUPABASE_URL", "https://valid.supabase.co")
    monkeypatch.setattr("app.core.config.settings.SUPABASE_SERVICE_ROLE_KEY", "secret_key_123")
    monkeypatch.setattr("app.core.config.settings.CV_STORAGE_BUCKET", "private_cv_bucket")

    user_id = uuid4()
    storage_path = f"{user_id}/my_cv.pdf"
    download_candidate_cv(user_id=user_id, storage_path=storage_path)

    mock_client.storage.from_.assert_called_once_with("private_cv_bucket")
    mock_bucket.download.assert_called_once_with(storage_path)


def test_storage_download_correct_user_owned_path(monkeypatch):
    """Test 15: Correct user-owned canonical storage path downloads successfully."""
    mock_client, _ = make_mock_supabase_client(b"file_bytes_123")
    monkeypatch.setattr("app.services.cv_storage.create_client", lambda url, key: mock_client)
    monkeypatch.setattr("app.core.config.settings.SUPABASE_URL", "https://valid.supabase.co")
    monkeypatch.setattr("app.core.config.settings.SUPABASE_SERVICE_ROLE_KEY", "secret_key_123")
    monkeypatch.setattr("app.core.config.settings.CV_STORAGE_BUCKET", "cvs")

    user_id = uuid4()
    res = download_candidate_cv(user_id=user_id, storage_path=f"{user_id}/resume.docx")
    assert res == b"file_bytes_123"


def test_cross_user_storage_path_rejected_before_network(monkeypatch):
    """
    Test 16: Cross-user storage paths are rejected before client creation.
    """
    client_calls = []
    monkeypatch.setattr(
        "app.services.cv_storage.create_client",
        lambda url, key: client_calls.append(1),
    )

    user_a = uuid4()
    user_b = uuid4()

    with pytest.raises(CVStorageValidationError, match="Unauthorized storage path access"):
        download_candidate_cv(user_id=user_a, storage_path=f"{user_b}/resume.pdf")

    assert client_calls == []


def test_unsupported_path_suffix_rejected_before_network(monkeypatch):
    """
    Test 17: Invalid storage suffixes are rejected before client creation.
    """
    client_calls = []
    monkeypatch.setattr(
        "app.services.cv_storage.create_client",
        lambda url, key: client_calls.append(1),
    )

    user_id = uuid4()
    with pytest.raises(CVStorageValidationError, match="Unsupported file extension"):
        download_candidate_cv(user_id=user_id, storage_path=f"{user_id}/malicious.exe")

    assert client_calls == []


def test_empty_provider_payload_rejected(monkeypatch):
    """
    Test 18: Empty Supabase Storage download payload is rejected.
    """
    mock_client, _ = make_mock_supabase_client(b"")
    monkeypatch.setattr("app.services.cv_storage.create_client", lambda url, key: mock_client)
    monkeypatch.setattr("app.core.config.settings.SUPABASE_URL", "https://valid.supabase.co")
    monkeypatch.setattr("app.core.config.settings.SUPABASE_SERVICE_ROLE_KEY", "secret_key_123")
    monkeypatch.setattr("app.core.config.settings.CV_STORAGE_BUCKET", "cvs")

    user_id = uuid4()
    with pytest.raises(CVStorageValidationError, match="Downloaded CV object is empty"):
        download_candidate_cv(user_id=user_id, storage_path=f"{user_id}/empty.pdf")


def test_download_provider_exception_propagates(monkeypatch):
    """Test 19: Supabase download exception propagates unchanged."""
    mock_client, mock_bucket = make_mock_supabase_client()
    mock_bucket.download.side_effect = RuntimeError("Supabase download timeout")
    monkeypatch.setattr("app.services.cv_storage.create_client", lambda url, key: mock_client)
    monkeypatch.setattr("app.core.config.settings.SUPABASE_URL", "https://valid.supabase.co")
    monkeypatch.setattr("app.core.config.settings.SUPABASE_SERVICE_ROLE_KEY", "secret_key_123")
    monkeypatch.setattr("app.core.config.settings.CV_STORAGE_BUCKET", "cvs")

    user_id = uuid4()
    with pytest.raises(RuntimeError, match="Supabase download timeout"):
        download_candidate_cv(user_id=user_id, storage_path=f"{user_id}/resume.pdf")


# ---------------------------------------------------------------------------
# 4. STRUCTURED PROFILE WRITE BOUNDARY TESTS (20 - 27)
# ---------------------------------------------------------------------------


def test_new_extracted_candidate_creates_profile_and_related_rows():
    """Test 20: replace_candidate_profile_from_extraction creates StudentProfile + child rows."""
    user_id = uuid4()
    db = TestingSessionLocal()
    try:
        extracted = _sample_extracted_profile()
        profile = replace_candidate_profile_from_extraction(
            db=db,
            user_id=user_id,
            cv_storage_path=f"{user_id}/resume.pdf",
            extracted=extracted,
        )
        db.commit()

        assert profile.full_name == "Samantha Ray"
        assert profile.headline == "AI Research Intern"

        # Verify child rows in DB
        skills = MatchingDataRepository.get_skill_names_for_student(db, profile.id)
        assert skills == ["PyTorch", "Python"]  # Alphabetical ordering

        education = MatchingDataRepository.get_education_for_student(db, profile.id)
        assert len(education) == 1
        assert education[0].institution == "MIT"

        experience = MatchingDataRepository.get_experience_for_student(db, profile.id)
        assert len(experience) == 1
        assert experience[0].company == "OpenLab"

        projects = MatchingDataRepository.get_projects_for_student(db, profile.id)
        assert len(projects) == 1
        assert projects[0].title == "Vision Transformer"
    finally:
        db.close()


def test_second_cv_extraction_replaces_prior_candidate_related_rows():
    """Test 21: Second CV extraction cleanly replaces all prior student child rows."""
    user_id = uuid4()
    db = TestingSessionLocal()
    try:
        # First extraction: Samantha Ray with MIT / OpenLab
        first_extracted = _sample_extracted_profile()
        prof1 = replace_candidate_profile_from_extraction(
            db=db,
            user_id=user_id,
            cv_storage_path=f"{user_id}/v1.pdf",
            extracted=first_extracted,
        )
        db.commit()

        # Second extraction: Updated Samantha with Harvard / Google
        second_extracted = ExtractedCandidateProfile(
            full_name="Samantha Ray",
            headline="Lead ML Engineer",
            skills=[ExtractedSkill(name="Rust"), ExtractedSkill(name="Go")],
            education=[
                ExtractedEducation(
                    institution="Harvard University",
                    degree="M.S. Computer Science",
                    start_year=2025,
                    end_year=2027,
                )
            ],
            experience=[
                ExtractedExperience(
                    company="Google DeepMind",
                    role="Research Scientist",
                    description="Advanced AI models.",
                )
            ],
            projects=[
                ExtractedProject(
                    title="Distributed Training",
                    tech_stack=["Rust", "Go"],
                )
            ],
            preferences=ExtractedPreferences(work_types=["remote"]),
        )

        prof2 = replace_candidate_profile_from_extraction(
            db=db,
            user_id=user_id,
            cv_storage_path=f"{user_id}/v2.pdf",
            extracted=second_extracted,
        )
        db.commit()

        assert prof2.id == prof1.id
        assert prof2.headline == "Lead ML Engineer"

        # Verify old MIT/OpenLab rows are gone and new Harvard/Google rows exist
        edu = MatchingDataRepository.get_education_for_student(db, prof2.id)
        assert len(edu) == 1
        assert edu[0].institution == "Harvard University"

        exp = MatchingDataRepository.get_experience_for_student(db, prof2.id)
        assert len(exp) == 1
        assert exp[0].company == "Google DeepMind"

        skills = MatchingDataRepository.get_skill_names_for_student(db, prof2.id)
        assert sorted(skills) == ["Go", "Rust"]
    finally:
        db.close()


def test_cv_extraction_replaces_skills_without_merging_old_skills():
    """
    Test 22: Subsequent CV extraction replaces candidate skills fresh
    and does not merge old ones.
    """
    user_id = uuid4()
    db = TestingSessionLocal()
    try:
        ext1 = ExtractedCandidateProfile(
            full_name="Dev",
            skills=[ExtractedSkill(name="Skill1"), ExtractedSkill(name="Skill2")],
        )
        prof = replace_candidate_profile_from_extraction(
            db=db, user_id=user_id, cv_storage_path="path", extracted=ext1
        )
        db.commit()

        # Update with Skill2 and Skill3 (Skill1 must be removed)
        ext2 = ExtractedCandidateProfile(
            full_name="Dev",
            skills=[ExtractedSkill(name="Skill2"), ExtractedSkill(name="Skill3")],
        )
        replace_candidate_profile_from_extraction(
            db=db, user_id=user_id, cv_storage_path="path", extracted=ext2
        )
        db.commit()

        student_skills = db.query(StudentSkill).filter_by(student_id=prof.id).all()
        assert len(student_skills) == 2
        skill_names = MatchingDataRepository.get_skill_names_for_student(db, prof.id)
        assert sorted(skill_names) == ["Skill2", "Skill3"]
    finally:
        db.close()


def test_global_skill_rows_not_deleted_on_student_skill_replacement():
    """Test 23: Global Skill taxonomy rows remain intact even when no student links to them."""
    user_id = uuid4()
    db = TestingSessionLocal()
    try:
        ext1 = ExtractedCandidateProfile(
            full_name="Candidate",
            skills=[ExtractedSkill(name="GlobalSkillA")],
        )
        replace_candidate_profile_from_extraction(
            db=db, user_id=user_id, cv_storage_path="p", extracted=ext1
        )
        db.commit()

        # Replace with GlobalSkillB
        ext2 = ExtractedCandidateProfile(
            full_name="Candidate",
            skills=[ExtractedSkill(name="GlobalSkillB")],
        )
        replace_candidate_profile_from_extraction(
            db=db, user_id=user_id, cv_storage_path="p", extracted=ext2
        )
        db.commit()

        # Global taxonomy must still contain GlobalSkillA
        skill_a = db.query(Skill).filter_by(name="GlobalSkillA").first()
        assert skill_a is not None
        skill_b = db.query(Skill).filter_by(name="GlobalSkillB").first()
        assert skill_b is not None
    finally:
        db.close()


def test_duplicate_extracted_skills_deduplicated_case_insensitively():
    """Test 24: Duplicate skills with different casing are deduplicated case-insensitively."""
    user_id = uuid4()
    db = TestingSessionLocal()
    try:
        extracted = ExtractedCandidateProfile(
            full_name="Test Dev",
            skills=[
                ExtractedSkill(name="Python"),
                ExtractedSkill(name="python"),
                ExtractedSkill(name="PYTHON"),
                ExtractedSkill(name="FastAPI"),
            ],
        )
        prof = replace_candidate_profile_from_extraction(
            db=db, user_id=user_id, cv_storage_path="p", extracted=extracted
        )
        db.commit()

        skills = MatchingDataRepository.get_skill_names_for_student(db, prof.id)
        assert skills == ["FastAPI", "Python"]
    finally:
        db.close()


def test_candidate_embedding_invalidated_before_regeneration():
    """Test 25: summary_embedding is set to None during structured profile write."""
    user_id = uuid4()
    db = TestingSessionLocal()
    try:
        # Create profile with an existing embedding
        prof = StudentProfileRepository.upsert_by_user_id(
            db=db, user_id=user_id, full_name="Initial Name"
        )
        StudentProfileRepository.set_summary_embedding(db, prof, [0.8] * 1536)
        db.commit()

        # Now replace candidate profile from extraction
        extracted = _sample_extracted_profile()
        updated_prof = replace_candidate_profile_from_extraction(
            db=db, user_id=user_id, cv_storage_path="p", extracted=extracted
        )
        assert updated_prof.summary_embedding is None
    finally:
        db.close()


def test_user_a_persistence_never_mutates_user_b():
    """
    Test 26: Replacing User A profile never mutates User B structured entities.
    """
    user_a = uuid4()
    user_b = uuid4()
    db = TestingSessionLocal()
    try:
        ext_b = ExtractedCandidateProfile(
            full_name="User B",
            headline="User B Headline",
            skills=[ExtractedSkill(name="UserBSkill")],
            education=[ExtractedEducation(institution="Uni B", degree="Deg B")],
        )
        prof_b = replace_candidate_profile_from_extraction(
            db=db, user_id=user_b, cv_storage_path="b.pdf", extracted=ext_b
        )
        db.commit()

        # Mutate User A
        ext_a = ExtractedCandidateProfile(
            full_name="User A",
            headline="User A Headline",
            skills=[ExtractedSkill(name="UserASkill")],
        )
        replace_candidate_profile_from_extraction(
            db=db, user_id=user_a, cv_storage_path="a.pdf", extracted=ext_a
        )
        db.commit()

        # Verify User B remained 100% untouched
        check_b = StudentProfileRepository.get_by_user_id(db, user_id=user_b)
        assert check_b.full_name == "User B"
        assert check_b.headline == "User B Headline"
        skills_b = MatchingDataRepository.get_skill_names_for_student(db, prof_b.id)
        assert skills_b == ["UserBSkill"]
        edu_b = MatchingDataRepository.get_education_for_student(db, prof_b.id)
        assert len(edu_b) == 1
        assert edu_b[0].institution == "Uni B"
    finally:
        db.close()


def test_repository_does_not_commit_or_rollback():
    """Test 27: replace_candidate_profile_from_extraction flushes but does not commit."""
    user_id = uuid4()
    db = TestingSessionLocal()
    try:
        extracted = _sample_extracted_profile()
        replace_candidate_profile_from_extraction(
            db=db, user_id=user_id, cv_storage_path="p", extracted=extracted
        )
        # Rollback without committing
        db.rollback()

        # Fresh session must find NO profile because repository never committed
        fresh_db = TestingSessionLocal()
        try:
            profile = StudentProfileRepository.get_by_user_id(fresh_db, user_id=user_id)
            assert profile is None
        finally:
            fresh_db.close()
    finally:
        db.close()


# ---------------------------------------------------------------------------
# 5. SEMANTIC VALIDATION TESTS (28 - 34)
# ---------------------------------------------------------------------------


def test_validate_cv_document_deterministic_sanity_empty_text():
    """Test 28: Empty or whitespace-only text raises InvalidCVDocumentError."""
    with pytest.raises(InvalidCVDocumentError, match="empty or contains only whitespace"):
        validate_cv_document("")

    with pytest.raises(InvalidCVDocumentError, match="empty or contains only whitespace"):
        validate_cv_document("   \n\t  ")


def test_validate_cv_document_deterministic_sanity_too_short_text():
    """Test 29: Insufficient text (< 40 chars or < 6 words) raises InvalidCVDocumentError."""
    with pytest.raises(InvalidCVDocumentError, match="insufficient text content"):
        validate_cv_document("Hello world")

    with pytest.raises(InvalidCVDocumentError, match="insufficient text content"):
        validate_cv_document("Short text 12345")


def test_validate_cv_document_plausible_cv_passes(monkeypatch):
    """Test 30: Plausible student CV text passes validation when semantic classifier approves."""
    mock_client = MagicMock()
    mock_response = MagicMock()
    mock_response.text = '{"is_cv": true, "confidence": 0.95, "reason_code": "valid_cv"}'
    mock_client.models.generate_content.return_value = mock_response

    monkeypatch.setattr("app.services.cv_validation.genai.Client", lambda api_key: mock_client)
    monkeypatch.setattr(
        "app.core.config.settings.GEMINI_API_KEY",
        "gemini-valid-test-key",
    )
    sample_cv = (
        "Jane Doe\n"
        "jane.doe@university.edu\n"
        "Education: Bachelor of Science in Computer Science, 2022-2026\n"
        "Skills: Python, FastAPI, React Native, Git, SQL\n"
        "Projects: InternMatch AI - Full-stack mobile platform"
    )
    result = validate_cv_document(sample_cv, content_locale="en")
    assert isinstance(result, CVValidationResult)
    assert result.is_cv is True
    assert result.confidence == 0.95
    assert result.reason_code == "valid_cv"


def test_validate_cv_document_unrelated_text_rejected_by_semantic_classifier(monkeypatch):
    """Test 31: Unrelated document (e.g. invoice/receipt) is rejected by classifier."""
    mock_client = MagicMock()
    mock_response = MagicMock()
    mock_response.text = (
        '{"is_cv": false, "confidence": 0.98, "reason_code": "invoice_or_financial"}'
    )
    mock_client.models.generate_content.return_value = mock_response

    monkeypatch.setattr("app.services.cv_validation.genai.Client", lambda api_key: mock_client)
    monkeypatch.setattr(
        "app.core.config.settings.GEMINI_API_KEY",
        "gemini-valid-test-key",
    )
    invoice_text = (
        "INVOICE #98765\n"
        "Billed To: ACME Corporation\n"
        "Item 1: Cloud Hosting Services - $450.00\n"
        "Item 2: Domain Registration - $50.00\n"
        "Total Balance Due: $500.00\n"
        "Please submit wire transfer to Bank Account 123456789."
    )
    with pytest.raises(InvalidCVDocumentError) as exc_info:
        validate_cv_document(invoice_text, content_locale="en")

    assert "does not appear to be a valid CV or resume" in str(exc_info.value)
    assert exc_info.value.reason_code == "invoice_or_financial"


def test_validate_cv_document_low_confidence_rejected(monkeypatch):
    """Test 32: Ambiguous document with low confidence (< 0.5) is rejected."""
    mock_client = MagicMock()
    mock_response = MagicMock()
    mock_response.text = '{"is_cv": true, "confidence": 0.40, "reason_code": "ambiguous_fragment"}'
    mock_client.models.generate_content.return_value = mock_response

    monkeypatch.setattr("app.services.cv_validation.genai.Client", lambda api_key: mock_client)
    monkeypatch.setattr(
        "app.core.config.settings.GEMINI_API_KEY",
        "gemini-valid-test-key",
    )
    text = "Some random text fragment discussing computer programming and software development."
    with pytest.raises(InvalidCVDocumentError):
        validate_cv_document(text, content_locale="en")


def test_validate_cv_document_provider_error_raises_service_error(monkeypatch):
    """Test 33: Provider exception raises CVValidationServiceError, not InvalidCVDocumentError."""
    mock_client = MagicMock()
    mock_client.models.generate_content.side_effect = RuntimeError("API connection timeout")

    monkeypatch.setattr("app.services.cv_validation.genai.Client", lambda api_key: mock_client)
    monkeypatch.setattr(
        "app.core.config.settings.GEMINI_API_KEY",
        "gemini-valid-test-key",
    )
    text = "Jane Doe Computer Science Student Resume Python SQL"
    with pytest.raises(CVValidationServiceError, match="LLM classification service error"):
        validate_cv_document(text, content_locale="en")


def test_validate_cv_document_turkish_and_arabic_supported(monkeypatch):
    """Test 34: Multilingual resumes (Turkish, Arabic) pass validation with target locale."""
    mock_client = MagicMock()
    mock_response = MagicMock()
    mock_response.text = '{"is_cv": true, "confidence": 0.96, "reason_code": "valid_cv"}'
    mock_client.models.generate_content.return_value = mock_response

    monkeypatch.setattr("app.services.cv_validation.genai.Client", lambda api_key: mock_client)
    monkeypatch.setattr(
        "app.core.config.settings.GEMINI_API_KEY",
        "gemini-valid-test-key",
    )
    turkish_cv = (
        "Ahmet Yılmaz\n"
        "Özgeçmiş\n"
        "Eğitim: Bilgisayar Mühendisliği, Boğaziçi Üniversitesi (2021-2025)\n"
        "Yetenekler: Python, Django, PostgreSQL, Docker\n"
        "Projeler: Akıllı Eşleştirme Uygulaması"
    )
    result = validate_cv_document(turkish_cv, content_locale="tr")
    assert result.is_cv is True
    assert result.confidence == 0.96


def test_validate_cv_document_multimodal_pdf_passes(monkeypatch):
    """Test 35: Multimodal PDF validation passes when semantic classifier approves bytes."""
    mock_client = MagicMock()
    mock_response = MagicMock()
    mock_response.text = '{"is_cv": true, "confidence": 0.97, "reason_code": "valid_cv"}'
    mock_client.models.generate_content.return_value = mock_response

    monkeypatch.setattr("app.services.cv_validation.genai.Client", lambda api_key: mock_client)
    monkeypatch.setattr("app.core.config.settings.GEMINI_API_KEY", "gemini-valid-test-key")

    pdf_bytes = b"%PDF-1.4 mock scanned image resume"
    result = validate_cv_document_multimodal(content=pdf_bytes, mime_type="application/pdf")

    assert isinstance(result, CVValidationResult)
    assert result.is_cv is True
    assert result.confidence == 0.97


def test_extract_structured_candidate_profile_multimodal_pdf(monkeypatch):
    """Test 36: Multimodal PDF structured extraction returns typed candidate profile."""
    mock_client = MagicMock()
    mock_response = MagicMock()
    mock_response.text = """{
        "full_name": "Visual Resume Candidate",
        "headline": "Full-Stack Designer",
        "skills": [
            {"name": "Figma", "proficiency_level": "advanced"},
            {"name": "React", "proficiency_level": "intermediate"}
        ],
        "education": [{
            "institution": "Design Academy",
            "degree": "B.A. Interaction Design",
            "start_year": 2021,
            "end_year": 2025
        }],
        "experience": [],
        "projects": [{
            "title": "Portfolio App",
            "tech_stack": ["React", "CSS"],
            "description": "Interactive showcase"
        }],
        "preferences": {
            "work_types": ["remote"],
            "desired_locations": ["Istanbul"],
            "target_roles": ["Product Designer"]
        }
    }"""
    mock_client.models.generate_content.return_value = mock_response

    monkeypatch.setattr(
        "app.services.cv_profile_extraction.genai.Client",
        lambda api_key: mock_client,
    )
    monkeypatch.setattr("app.core.config.settings.GEMINI_API_KEY", "gemini-valid-test-key")

    pdf_bytes = b"%PDF-1.4 mock complex layout resume"
    extracted = extract_structured_candidate_profile_multimodal(
        content=pdf_bytes,
        mime_type="application/pdf",
    )

    assert isinstance(extracted, ExtractedCandidateProfile)
    assert extracted.full_name == "Visual Resume Candidate"
    assert extracted.headline == "Full-Stack Designer"
    assert len(extracted.skills) == 2
    assert extracted.skills[0].name == "Figma"
    assert len(extracted.education) == 1
    assert extracted.education[0].institution == "Design Academy"
